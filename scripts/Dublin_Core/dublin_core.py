import PyPDF2
import openpyxl
import json
import hashlib
import mimetypes
from rdflib import Graph, Namespace, Literal, URIRef
from rdflib.namespace import RDF, RDFS, XSD, DC, DCTERMS
from datetime import datetime
from pathlib import Path
import re
import os
import csv
import sys
from typing import Dict, List, Optional, Any
from dotenv import load_dotenv
from Dublin_Core.dublin_core_hana import store_dublin_core_metadata


# New imports for DOCX/DOC support
from docx import Document
import zipfile
from xml.etree import ElementTree as ET

# Load environment variables
load_dotenv()

# Define namespaces
DCO = Namespace("http://example.org/doc-copyright-ontology#")
FOAF = Namespace("http://xmlns.com/foaf/0.1/")
CC = Namespace("http://creativecommons.org/ns#")

class EnhancedDocumentMetadataExtractor:
    def __init__(self):
        self.graph = Graph()
        # Bind namespaces
        self.graph.bind("dco", DCO)
        self.graph.bind("dc", DC)
        self.graph.bind("dcterms", DCTERMS)
        self.graph.bind("foaf", FOAF)
        self.graph.bind("cc", CC)
        
        # Define supported formats - now includes DOCX and DOC
        self.supported_formats = {
            '.pdf': 'PDF',
            '.xlsx': 'XLSX',
            '.docx': 'DOCX',
            '.doc': 'DOC',
            '.jpg': 'JPG',
            '.jpeg': 'JPEG',
            '.png': 'PNG'
        }
        
        # Define document types
        self.document_types = {
            'pdf': 'Text',
            'xlsx': 'Dataset',
            'docx': 'Text',
            'doc': 'Text',
            'jpg': 'Image',
            'jpeg': 'Image',
            'png': 'Image'
        }

    def calculate_file_hash(self, file_path: str) -> Dict[str, str]:
        """Calculate MD5 and SHA256 hashes for the file"""
        md5_hash = hashlib.md5()
        sha256_hash = hashlib.sha256()
        
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                md5_hash.update(chunk)
                sha256_hash.update(chunk)
        
        return {
            'md5': md5_hash.hexdigest(),
            'sha256': sha256_hash.hexdigest()
        }

    def get_file_info(self, file_path: str) -> Dict[str, Any]:
        """Get basic file information"""
        path_obj = Path(file_path)
        file_stats = path_obj.stat()
        
        return {
            'file_name': path_obj.name,
            'file_path': str(path_obj.parent),
            'file_size': file_stats.st_size,
            'file_hash': self.calculate_file_hash(file_path),
            'mime_type': mimetypes.guess_type(file_path)[0] or 'application/octet-stream'
        }

    def scan_document_for_copyright(self, file_path: str) -> Dict[str, Any]:
        """Enhanced copyright scanning for PDF, XLSX, DOCX, and DOC files"""
        file_ext = Path(file_path).suffix.lower()
        copyright_info = {}
        
        try:
            if file_ext == '.pdf':
                copyright_info = self._scan_pdf_for_copyright(file_path)
            elif file_ext == '.xlsx':
                copyright_info = self._scan_xlsx_for_copyright(file_path)
            elif file_ext == '.docx':
                copyright_info = self._scan_docx_for_copyright(file_path)
            elif file_ext == '.doc':
                copyright_info = self._scan_doc_for_copyright(file_path)
            else:
                # For image files, use basic file metadata
                copyright_info = self._get_default_copyright_info(file_path)
                
        except Exception as e:
            print(f"Error scanning document for copyright: {e}")
            copyright_info = self._get_default_copyright_info(file_path)
        
        return copyright_info

    def _scan_pdf_for_copyright(self, pdf_path: str) -> Dict[str, Any]:
        """Scan PDF document thoroughly for copyright information"""
        copyright_info = {}
        
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            # Extract text from all pages for comprehensive copyright scanning
            full_text = ""
            max_pages = min(len(pdf_reader.pages), 10)  # Scan first 10 pages
            
            for page_num in range(max_pages):
                page = pdf_reader.pages[page_num]
                full_text += page.extract_text() + "\n"
            
            # Enhanced copyright extraction
            copyright_info = self._extract_enhanced_copyright_info(full_text)
            
            # Also check PDF metadata for additional copyright info
            if pdf_reader.metadata:
                metadata = pdf_reader.metadata
                if '/Copyright' in metadata:
                    copyright_info['metadata_copyright'] = metadata['/Copyright']
                if '/Rights' in metadata:
                    copyright_info['metadata_rights'] = metadata['/Rights']
        
        return copyright_info

    def _scan_xlsx_for_copyright(self, xlsx_path: str) -> Dict[str, Any]:
        """Scan XLSX document thoroughly for copyright information"""
        copyright_info = {}
        
        workbook = openpyxl.load_workbook(xlsx_path, read_only=True, data_only=True)
        
        # Extract text from all worksheets
        full_text = ""
        
        for sheet in workbook.worksheets:
            # Scan more rows for comprehensive copyright detection
            max_rows = min(sheet.max_row, 50) if sheet.max_row else 50
            
            for row in sheet.iter_rows(max_row=max_rows, values_only=True):
                for cell in row:
                    if cell and isinstance(cell, str):
                        full_text += str(cell) + " "
        
        # Enhanced copyright extraction
        copyright_info = self._extract_enhanced_copyright_info(full_text)
        
        # Check workbook properties for additional copyright info
        props = workbook.properties
        if props.description:
            desc_copyright = self._extract_enhanced_copyright_info(props.description)
            copyright_info.update(desc_copyright)
        
        workbook.close()
        
        return copyright_info

    def _scan_docx_for_copyright(self, docx_path: str) -> Dict[str, Any]:
        """Scan DOCX document thoroughly for copyright information"""
        copyright_info = {}
        
        try:
            doc = Document(docx_path)
            
            # Extract text from all paragraphs
            full_text = ""
            for paragraph in doc.paragraphs:
                full_text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text += cell.text + " "
            
            # Extract text from headers and footers
            for section in doc.sections:
                # Header text
                if section.header:
                    for paragraph in section.header.paragraphs:
                        full_text += paragraph.text + "\n"
                
                # Footer text
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        full_text += paragraph.text + "\n"
            
            # Enhanced copyright extraction
            copyright_info = self._extract_enhanced_copyright_info(full_text)
            
            # Check document properties for additional copyright info
            core_props = doc.core_properties
            if hasattr(core_props, 'comments') and core_props.comments:
                desc_copyright = self._extract_enhanced_copyright_info(core_props.comments)
                copyright_info.update(desc_copyright)
            
            if hasattr(core_props, 'description') and core_props.description:
                desc_copyright = self._extract_enhanced_copyright_info(core_props.description)
                copyright_info.update(desc_copyright)
                
        except Exception as e:
            print(f"Error scanning DOCX file: {e}")
            copyright_info = self._get_default_copyright_info(docx_path)
        
        return copyright_info

    def _scan_doc_for_copyright(self, doc_path: str) -> Dict[str, Any]:
        """Scan DOC document for copyright information (limited functionality)"""
        copyright_info = {}
        
        try:
            # For .doc files, we'll use a basic approach since python-docx doesn't support them
            # This is a fallback method - in production, you might want to use python-docx2txt
            # or convert to DOCX format first
            
            # Try to extract some basic text using a simple binary search approach
            # This is not ideal but provides some functionality for .doc files
            with open(doc_path, 'rb') as file:
                content = file.read()
                
                # Convert bytes to string, ignoring errors
                try:
                    text_content = content.decode('utf-8', errors='ignore')
                except:
                    text_content = content.decode('latin-1', errors='ignore')
                
                # Extract copyright information from whatever text we can find
                copyright_info = self._extract_enhanced_copyright_info(text_content)
                
        except Exception as e:
            print(f"Error scanning DOC file (limited support): {e}")
            copyright_info = self._get_default_copyright_info(doc_path)
        
        return copyright_info

    def _extract_enhanced_copyright_info(self, text: str) -> Dict[str, Any]:
        """Enhanced copyright information extraction with better pattern matching"""
        copyright_info = {}
        
        # Multiple copyright year patterns
        copyright_patterns = [
            r'(?:©|Copyright|copyright|COPYRIGHT)\s*(\d{4}(?:-\d{4})?)\s*(?:by\s+)?([^.\n\r,;]+)',
            r'(?:©|Copyright|copyright|COPYRIGHT)\s*([^.\n\r,;]+)\s*(\d{4}(?:-\d{4})?)',
            r'(?:©|Copyright|copyright|COPYRIGHT)\s*(\d{4}(?:-\d{4})?)',
            r'(?:©|Copyright|copyright|COPYRIGHT)\s*([^.\n\r,;]+)',
            r'All rights reserved\s*(?:©|Copyright|copyright|COPYRIGHT)?\s*(\d{4}(?:-\d{4})?)\s*([^.\n\r,;]+)?'
        ]
        
        for pattern in copyright_patterns:
            copyright_match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if copyright_match:
                groups = copyright_match.groups()
                
                # Determine which group contains year and which contains holder
                for group in groups:
                    if group and re.match(r'\d{4}(?:-\d{4})?', group.strip()):
                        year_match = re.search(r'(\d{4})', group)
                        if year_match:
                            copyright_info['copyright_year'] = int(year_match.group(1))
                    elif group and len(group.strip()) > 0:
                        holder = group.strip()
                        # Clean up common suffixes/prefixes
                        holder = re.sub(r'^(?:by\s+|©\s*|copyright\s*)', '', holder, flags=re.IGNORECASE)
                        holder = re.sub(r'(?:\s*all rights reserved|\s*\.|\s*,)$', '', holder, flags=re.IGNORECASE)
                        if len(holder) > 0:
                            copyright_info['copyright_holder'] = holder
                
                break
        
        # Enhanced license and rights detection
        license_patterns = [
            (r'Creative Commons\s+(?:Attribution\s+)?(?:Share\s*Alike\s+)?(?:Non\s*Commercial\s+)?[\d\.]*\s*(?:CC\s+BY(?:-NC)?(?:-SA)?)?', 'Creative Commons'),
            (r'CC\s+BY(?:-NC)?(?:-SA)?(?:-ND)?(?:\s+[\d\.]+)?', 'Creative Commons'),
            (r'MIT\s+License', 'MIT'),
            (r'GNU\s+General\s+Public\s+License|GPL', 'GPL'),
            (r'Apache\s+License', 'Apache'),
            (r'BSD\s+License', 'BSD'),
            (r'Public\s+Domain', 'Public Domain'),
            (r'All\s+rights\s+reserved', 'All Rights Reserved'),
            (r'No\s+rights\s+reserved', 'Public Domain'),
            (r'Proprietary', 'Proprietary')
        ]
        
        for pattern, license_type in license_patterns:
            license_match = re.search(pattern, text, re.IGNORECASE)
            if license_match:
                copyright_info['license_type'] = license_type
                copyright_info['rights'] = license_match.group(0)
                
                # Determine copyright status based on license
                if license_type == 'Public Domain':
                    copyright_info['copyright_status'] = 'Public Domain'
                elif license_type in ['Creative Commons', 'MIT', 'GPL', 'Apache', 'BSD']:
                    copyright_info['copyright_status'] = 'Open Source'
                else:
                    copyright_info['copyright_status'] = 'Copyright'
                break
        
        # Additional copyright status indicators
        if 'copyright_status' not in copyright_info:
            status_patterns = [
                (r'(?:©|Copyright|copyright)', 'Copyright'),
                (r'All\s+rights\s+reserved', 'Copyright'),
                (r'Proprietary', 'Copyright'),
                (r'Confidential', 'Copyright')
            ]
            
            for pattern, status in status_patterns:
                if re.search(pattern, text, re.IGNORECASE):
                    copyright_info['copyright_status'] = status
                    break
        
        # Set defaults if not found
        if 'copyright_status' not in copyright_info:
            copyright_info['copyright_status'] = 'Unknown'
        
        if 'license_type' not in copyright_info:
            copyright_info['license_type'] = 'All Rights Reserved'
        
        return copyright_info

    def _get_default_copyright_info(self, file_path: str) -> Dict[str, Any]:
        """Get default copyright information for files without scannable content"""
        return {
            'copyright_status': 'Unknown',
            'license_type': 'All Rights Reserved'
        }

    def extract_pdf_metadata(self, pdf_path: str) -> Dict[str, Any]:
        """Extract metadata from PDF file"""
        metadata = {'format': 'PDF', 'type': 'Text'}
        
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                if pdf_reader.metadata:
                    info = pdf_reader.metadata
                    metadata['title'] = info.get('/Title', 'Untitled')
                    metadata['creator'] = [info.get('/Author', 'Unknown')] if info.get('/Author') else ['Unknown']
                    metadata['subject'] = info.get('/Subject', '').split(',') if info.get('/Subject') else []
                    
                    if '/CreationDate' in info:
                        creation_date = info['/CreationDate']
                        if creation_date.startswith('D:'):
                            try:
                                date_str = creation_date[2:16]
                                metadata['date'] = datetime.strptime(date_str, '%Y%m%d%H%M%S').isoformat() + 'Z'
                            except ValueError:
                                metadata['date'] = datetime.now().isoformat() + 'Z'
                
                # Enhanced copyright extraction using the new scanning method
                copyright_info = self.scan_document_for_copyright(pdf_path)
                metadata.update(copyright_info)
                    
        except Exception as e:
            print(f"Error extracting PDF metadata: {e}")
            
        return metadata

    def extract_excel_metadata(self, excel_path: str) -> Dict[str, Any]:
        """Extract metadata from Excel file"""
        file_ext = Path(excel_path).suffix.upper().replace('.', '')
        metadata = {'format': file_ext, 'type': 'Dataset'}
        
        try:
            workbook = openpyxl.load_workbook(excel_path, read_only=True, data_only=True)
            
            props = workbook.properties
            metadata['title'] = props.title or 'Untitled'
            metadata['creator'] = [props.creator] if props.creator else ['Unknown']
            metadata['subject'] = props.subject.split(',') if props.subject else []
            metadata['description'] = props.description or ''
            metadata['date'] = props.created.isoformat() + 'Z' if props.created else datetime.now().isoformat() + 'Z'
            
            # Enhanced copyright extraction using the new scanning method
            copyright_info = self.scan_document_for_copyright(excel_path)
            metadata.update(copyright_info)
                    
            workbook.close()
            
        except Exception as e:
            print(f"Error extracting Excel metadata: {e}")
            
        return metadata

    def extract_docx_metadata(self, docx_path: str) -> Dict[str, Any]:
        """Extract metadata from DOCX file"""
        metadata = {'format': 'DOCX', 'type': 'Text'}
        
        try:
            doc = Document(docx_path)
            core_props = doc.core_properties
            
            metadata['title'] = core_props.title or 'Untitled'
            metadata['creator'] = [core_props.author] if core_props.author else ['Unknown']
            metadata['subject'] = core_props.subject.split(',') if core_props.subject else []
            metadata['description'] = core_props.comments or ''
            metadata['date'] = core_props.created.isoformat() + 'Z' if core_props.created else datetime.now().isoformat() + 'Z'
            
            # Add additional metadata if available
            if core_props.last_modified_by:
                metadata['contributor'] = [core_props.last_modified_by]
            else:
                metadata['contributor'] = []
            
            if core_props.keywords:
                # Add keywords to subject if not already present
                keywords = [kw.strip() for kw in core_props.keywords.split(',') if kw.strip()]
                metadata['subject'].extend(keywords)
            
            # Enhanced copyright extraction using the new scanning method
            copyright_info = self.scan_document_for_copyright(docx_path)
            metadata.update(copyright_info)
                    
        except Exception as e:
            print(f"Error extracting DOCX metadata: {e}")
            
        return metadata

    def extract_doc_metadata(self, doc_path: str) -> Dict[str, Any]:
        """Extract metadata from DOC file (limited functionality)"""
        metadata = {'format': 'DOC', 'type': 'Text'}
        
        try:
            # For .doc files, we have limited metadata extraction capabilities
            # Use file system metadata as fallback
            metadata['title'] = Path(doc_path).stem
            metadata['creator'] = ['Unknown']
            metadata['subject'] = []
            metadata['description'] = ''
            metadata['date'] = datetime.fromtimestamp(Path(doc_path).stat().st_mtime).isoformat() + 'Z'
            metadata['contributor'] = []
            
            # Enhanced copyright extraction using the new scanning method
            copyright_info = self.scan_document_for_copyright(doc_path)
            metadata.update(copyright_info)
                    
        except Exception as e:
            print(f"Error extracting DOC metadata: {e}")
            
        return metadata

    def extract_image_metadata(self, image_path: str) -> Dict[str, Any]:
        """Extract metadata from image files"""
        file_ext = Path(image_path).suffix.upper().replace('.', '')
        metadata = {'format': file_ext, 'type': 'Image'}
        
        try:
            metadata['title'] = Path(image_path).stem
            metadata['creator'] = ['Unknown']
            metadata['date'] = datetime.fromtimestamp(Path(image_path).stat().st_mtime).isoformat() + 'Z'
            
            # Use basic copyright info for images
            copyright_info = self.scan_document_for_copyright(image_path)
            metadata.update(copyright_info)
            
        except Exception as e:
            print(f"Error extracting image metadata: {e}")
            
        return metadata

    def extract_copyright_from_text(self, text: str) -> Dict[str, Any]:
        """Extract copyright information from text content (legacy method - now uses enhanced scanning)"""
        return self._extract_enhanced_copyright_info(text)

    def extract_metadata_by_format(self, file_path: str) -> Dict[str, Any]:
        """Extract metadata based on file format"""
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.pdf':
            return self.extract_pdf_metadata(file_path)
        elif file_ext == '.xlsx':
            return self.extract_excel_metadata(file_path)
        elif file_ext == '.docx':
            return self.extract_docx_metadata(file_path)
        elif file_ext == '.doc':
            return self.extract_doc_metadata(file_path)
        elif file_ext in ['.jpg', '.jpeg', '.png']:
            return self.extract_image_metadata(file_path)
        else:
            # Default metadata for unsupported formats
            return {
                'title': Path(file_path).stem,
                'creator': ['Unknown'],
                'date': datetime.fromtimestamp(Path(file_path).stat().st_mtime).isoformat() + 'Z',
                'format': file_ext.upper().replace('.', ''),
                'type': 'PhysicalObject',
                'copyright_status': 'Unknown',
                'license_type': 'All Rights Reserved'
            }

    def _determine_copyright_holder_type(self, holder_name: str) -> str:
        """Determine if copyright holder is an Organization or Person based on name patterns"""
        if not holder_name:
            return "Unknown"
        
        # Common organizational indicators
        org_indicators = [
            'corp', 'inc', 'ltd', 'llc', 'company', 'corporation', 'limited', 'enterprises',
            'group', 'holdings', 'international', 'global', 'worldwide', 'publishing',
            'press', 'media', 'news', 'times', 'post', 'journal', 'magazine', 'books',
            'university', 'college', 'institute', 'foundation', 'association', 'society',
            'council', 'board', 'committee', 'department', 'ministry', 'agency', 'bureau',
            'government', 'federal', 'state', 'national', 'public', 'private', 'llp',
            'partnership', 'trust', 'fund', 'bank', 'financial', 'insurance', 'consulting',
            'services', 'solutions', 'technologies', 'systems', 'software', 'networks'
        ]
        
        holder_lower = holder_name.lower()
        
        # Check for organizational indicators
        for indicator in org_indicators:
            if indicator in holder_lower:
                return "Organization"
        
        # Check for multiple words (often indicates organization)
        words = holder_name.split()
        if len(words) > 2:
            return "Organization"
        
        # Check for typical personal name patterns (First Last, First M. Last, etc.)
        if len(words) == 2 and all(word.isalpha() or '.' in word for word in words):
            return "Person"
        
        # Default to Organization if uncertain
        return "Organization"

    def create_dublin_core_json(self, file_path: str, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Create Dublin Core JSON following your schema with string copyright holder"""
        file_info = self.get_file_info(file_path)
        doc_id = f"DOC-{datetime.now().strftime('%Y')}-{abs(hash(file_path)) % 1000000:06d}"
        
        # Create copyright holder string
        copyright_holder_str = ""
        if metadata.get('copyright_holder'):
            holder_name = metadata['copyright_holder']
            holder_type = self._determine_copyright_holder_type(holder_name)
            copyright_holder_str = f"{holder_name} ({holder_type})"
        
        dublin_core_json = {
            "document_id": doc_id,
            "dublin_core": {
                "title": metadata.get('title', 'Untitled'),
                "creator": metadata.get('creator', ['Unknown']),
                "subject": metadata.get('subject', []),
                "description": metadata.get('description', ''),
                "publisher": metadata.get('publisher', ''),
                "contributor": metadata.get('contributor', []),
                "date": metadata.get('date', datetime.now().isoformat() + 'Z'),
                "type": metadata.get('type', 'Text'),
                "format": metadata.get('format', 'Unknown'),
                "language": metadata.get('language', 'en-US'),
                "rights": metadata.get('rights', '')
            },
            "copyright_info": {
                "copyright_year": metadata.get('copyright_year'),
                "copyright_holder": copyright_holder_str,
                "copyright_status": metadata.get('copyright_status', 'Unknown')
            },
            "license_info": {
                "license_type": metadata.get('license_type', 'All Rights Reserved')
            },
            "access_control": {
                "access_level": "Public"
            },
            "extraction_metadata": {
                "extraction_method": "LLM",
                "extraction_date": datetime.now().isoformat() + 'Z',
                "extraction_tool": "Enhanced Dublin Core Extractor v3.0",
                "confidence_scores": {
                    "overall": 0.85,
                    "title": 0.95,
                    "creator": 0.80,
                    "copyright_info": 0.75,
                    "license_info": 0.70
                }
            },
            "file_info": file_info,
            "audit_trail": [
                {
                    "timestamp": datetime.now().isoformat() + 'Z',
                    "user": "system",
                    "action": "Created",
                    "changes": {"initial_extraction": True}
                }
            ]
        }
        
        return dublin_core_json

    def process_file(self, file_path: str, output_dir: str = None) -> Optional[Dict[str, Any]]:
        """Process a single file and generate both JSON and TTL outputs"""
        if not os.path.exists(file_path):
            print(f"Error: File {file_path} not found")
            return None
        
        file_ext = Path(file_path).suffix.lower()
        if file_ext not in self.supported_formats:
            print(f"Error: File type {file_ext} not supported. Supported formats: {list(self.supported_formats.keys())}")
            return None
        
        print(f"Processing file: {Path(file_path).name}")
        print(f"Scanning document for copyright information...")
        
        # Extract metadata
        metadata = self.extract_metadata_by_format(file_path)
        
        # Create Dublin Core JSON
        dublin_core_json = self.create_dublin_core_json(file_path, metadata)
        
        # Store metadata in HANA
        try:
            store_dublin_core_metadata(dublin_core_json, Path(file_path).name)
        except Exception as e:
            print(f"Error storing metadata in HANA: {str(e)}")
        
        # Set output directory
        if output_dir is None:
            output_dir = Path(file_path).parent
        
        # Save JSON output
        base_name = Path(file_path).stem
        json_path = Path(output_dir) / f"{base_name}_metadata.json"
        
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(dublin_core_json, f, indent=2, ensure_ascii=False)
        
        # Create RDF representation
        doc_uri = self.create_rdf_document(file_path, dublin_core_json)
        
        # Save TTL output
        ttl_path = Path(output_dir) / f"{base_name}_metadata.ttl"
        self.export_to_turtle(str(ttl_path))
        
        print(f"Generated metadata files:")
        print(f"  - JSON: {json_path}")
        print(f"  - TTL: {ttl_path}")
        
        # Print copyright information found
        copyright_info = dublin_core_json.get('copyright_info', {})
        license_info = dublin_core_json.get('license_info', {})
        
        print(f"\nCopyright Information Extracted:")
        print(f"  - Copyright Year: {copyright_info.get('copyright_year', 'Not found')}")
        print(f"  - Copyright Holder: {copyright_info.get('copyright_holder', 'Not found')}")
        print(f"  - Copyright Status: {copyright_info.get('copyright_status', 'Unknown')}")
        print(f"  - License Type: {license_info.get('license_type', 'All Rights Reserved')}")
        
        # Clear graph for next file
        self.graph = Graph()
        self.graph.bind("dco", DCO)
        self.graph.bind("dc", DC)
        self.graph.bind("dcterms", DCTERMS)
        self.graph.bind("foaf", FOAF)
        self.graph.bind("cc", CC)
        
        return {
            'file_path': file_path,
            'json_output': str(json_path),
            'ttl_output': str(ttl_path),
            'metadata': dublin_core_json
        }

    def create_rdf_document(self, file_path: str, dublin_core_json: Dict[str, Any]) -> URIRef:
        """Create RDF representation from Dublin Core JSON"""
        doc_uri = URIRef(f"http://example.org/documents/{dublin_core_json['document_id']}")
        
        # Add Dublin Core metadata
        dc = dublin_core_json['dublin_core']
        self.graph.add((doc_uri, DCO.title, Literal(dc['title'])))
        
        for creator in dc['creator']:
            self.graph.add((doc_uri, DCO.creator, Literal(creator)))
        
        self.graph.add((doc_uri, DCO.date, Literal(dc['date'], datatype=XSD.dateTime)))
        
        # Fix: Use proper RDF term instead of DCO.format which conflicts with Python's format method
        self.graph.add((doc_uri, DCTERMS.format, Literal(dc['format'])))
        
        self.graph.add((doc_uri, DCO.type, Literal(dc['type'])))
        self.graph.add((doc_uri, DCTERMS.description, Literal(dc['description'])))
        self.graph.add((doc_uri, DCTERMS.language, Literal(dc['language'])))
        self.graph.add((doc_uri, DCTERMS.rights, Literal(dc['rights'])))
        
        # Add subjects
        for subject in dc['subject']:
            if subject:  # Only add non-empty subjects
                self.graph.add((doc_uri, DCTERMS.subject, Literal(subject)))
        
        # Add contributors
        for contributor in dc['contributor']:
            if contributor:  # Only add non-empty contributors
                self.graph.add((doc_uri, DCTERMS.contributor, Literal(contributor)))
        
        # Add publisher if present
        if dc['publisher']:
            self.graph.add((doc_uri, DCTERMS.publisher, Literal(dc['publisher'])))
        
        # Add copyright information - updated for string copyright holder
        copyright_info = dublin_core_json.get('copyright_info', {})
        if copyright_info.get('copyright_year'):
            self.graph.add((doc_uri, DCO.copyrightYear, Literal(copyright_info['copyright_year'], datatype=XSD.gYear)))
        
        if copyright_info.get('copyright_status'):
            self.graph.add((doc_uri, DCO.copyrightStatus, Literal(copyright_info['copyright_status'])))
        
        # Add copyright holder as string
        if copyright_info.get('copyright_holder'):
            self.graph.add((doc_uri, DCO.copyrightHolder, Literal(copyright_info['copyright_holder'])))
        
        # Add license information
        license_info = dublin_core_json.get('license_info', {})
        if license_info.get('license_type'):
            self.graph.add((doc_uri, DCO.licenseType, Literal(license_info['license_type'])))
        
        # Add file information
        file_info = dublin_core_json.get('file_info', {})
        if file_info.get('file_size'):
            self.graph.add((doc_uri, DCO.fileSize, Literal(file_info['file_size'], datatype=XSD.integer)))
        
        if file_info.get('mime_type'):
            self.graph.add((doc_uri, DCO.mimeType, Literal(file_info['mime_type'])))
        
        # Add file hashes
        file_hash = file_info.get('file_hash', {})
        if file_hash.get('md5'):
            self.graph.add((doc_uri, DCO.md5Hash, Literal(file_hash['md5'])))
        if file_hash.get('sha256'):
            self.graph.add((doc_uri, DCO.sha256Hash, Literal(file_hash['sha256'])))
        
        return doc_uri

    def scan_directory(self, directory_path: str, output_dir: str = None) -> List[Dict[str, Any]]:
        """Scan directory and process all supported files"""
        directory = Path(directory_path)
        if not directory.exists():
            print(f"Error: Directory {directory_path} not found")
            return []
        
        processed_files = []
        supported_extensions = list(self.supported_formats.keys())
        
        for file_path in directory.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                try:
                    result = self.process_file(str(file_path), output_dir)
                    if result:
                        processed_files.append(result)
                except Exception as e:
                    print(f"Error processing {file_path}: {e}")
        
        return processed_files

    def export_to_turtle(self, output_file: str):
        """Export the RDF graph to Turtle format"""
        self.graph.serialize(destination=output_file, format='turtle')

    def process_single_file_by_name(self, filename: str, search_path: str = None, output_dir: str = None) -> Optional[Dict[str, Any]]:
        """Process a single file by filename, optionally searching in a specific directory"""
        if search_path is None:
            search_path = os.getenv('LOCALPATH', '') + "Documents"
        
        search_directory = Path(search_path)
        
        # If filename is already a full path, use it directly
        if Path(filename).is_absolute() and Path(filename).exists():
            return self.process_file(filename, output_dir)
        
        # Search for the file in the specified directory
        found_files = []
        for file_path in search_directory.rglob(filename):
            if file_path.is_file():
                found_files.append(file_path)
        
        if not found_files:
            print(f"Error: File '{filename}' not found in directory '{search_path}'")
            return None
        
        if len(found_files) > 1:
            print(f"Warning: Multiple files found with name '{filename}':")
            for i, file_path in enumerate(found_files, 1):
                print(f"  {i}. {file_path}")
            print(f"Processing the first one: {found_files[0]}")
        
        # Process the first (or only) found file
        return self.process_file(str(found_files[0]), output_dir)